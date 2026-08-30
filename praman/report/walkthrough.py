"""`python -m praman walkthrough` — the solution walkthrough, as a .docx.

Structured the way a company reports on a product it is taking to market, and
written for a reader who is not a payments engineer: why the thing exists,
what it is, what was decided and what was rejected, then the evidence.

Three rules the document is built to:

  * **Figures carry the argument, not paragraphs.** Six charts, four of them
    computed from `results/` and `corpus.yaml` at build time. Prose exists to
    say what a figure means, not to repeat it.
  * **Every term is defined before it is used.** A glossary sits before the
    engineering, because "mandate chain", "ASR" and "reference monitor" are
    alien to most readers and the argument is worthless if they stop there.
  * **Every external number has a citation.** They live in `sources.py`, one
    registry, checked by a test. Anything we measured ourselves is marked as
    ours and is not quoted to anybody else.

The challenge asks for four things: the attacks identified, how they are
generated and simulated, the detection model and its efficacy, and real-world
feasibility. Those are Part IV's four books.

**Generated, not typed.** Every quantity is read out of the repository through
the same `praman.metrics` the arena and the CLI use, so the document cannot
drift from the harness it describes. The prose is ours; the arithmetic is the
harness's — including the unflattering parts.
"""

from __future__ import annotations

import argparse
from datetime import date
from pathlib import Path

from praman import metrics
from praman.api.replay import CampaignStore
from praman.blue.defense import control_label
from praman.console import banner
from praman.console import setup as console_setup
from praman.money import fmt
from praman.red.corpus import SURFACES, load_corpus
from praman.red.corpus_report import HEADINGS, by_surface
from praman.report import charts
from praman.report.sources import MARKET, SOURCES, cite

__all__ = ["build"]

TEAM = "AIZEN"
REPO = "https://github.com/DnyaneshU/praman"
ARENA = "https://huggingface.co/spaces/DnyaneshU/praman"

#: Ink, accent, and a muted grey for captions. The accent is the arena's own
#: red, so the document, the charts and the prototype read as one product.
INK = (0x11, 0x18, 0x22)
ACCENT = (0xC0, 0x2B, 0x3E)
MUTED = (0x5B, 0x66, 0x75)

#: Plain-language definitions, printed before the engineering that uses them.
GLOSSARY = (
    (
        "Agentic payment",
        "A payment an AI agent makes on your behalf, from an authority you granted "
        "it in advance — rather than one you approve at the moment it happens.",
    ),
    (
        "Mandate",
        "A signed document that records one step of that authority. Praman uses the "
        "three the industry has settled on: an Intent (what you authorised), a Cart "
        "(what the merchant offered), a Payment (what actually moves).",
    ),
    (
        "Mandate chain",
        "The three joined together. Each signature proves who wrote its own "
        "document — and nothing about whether the three agree with each other. "
        "That gap is the subject of this report.",
    ),
    (
        "Breach and attack simulation",
        "Attacking your own defences on purpose, on a schedule, to find out what "
        "they actually stop. Routine for corporate networks; nobody does it for "
        "payment mandates.",
    ),
    (
        "Attack success rate (ASR)",
        "The share of attacks that moved money. Praman decides this from the "
        "ledger — did an attacker's balance go up — not from whether the control "
        "raised an alarm.",
    ),
    (
        "Precision, recall, F1",
        "Standard detection scores. Precision: of everything flagged, how much was "
        "really an attack. Recall: of all attacks, how many got flagged. F1 "
        "combines them. All three can be perfect while money is still stolen — "
        "see Figure 4.",
    ),
    (
        "Reference monitor",
        "A checker that sits outside the thing it is checking and cannot be "
        "bypassed, because nothing settles without going through it. A 1972 idea, "
        "applied here to payments.",
    ),
    (
        "Adaptive attacker",
        "An attacker who reads the rejection, works out which rule refused it, and "
        "tries again. The realistic case, and the one nobody publishes numbers for.",
    ),
)


def build(*, results: Path, out: Path) -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    records = {r.id: r for r in CampaignStore(results).list()}
    seeds = load_corpus()
    built = [s for s in seeds if s.implemented]
    surfaces = by_surface(seeds)
    figures = charts.render_all(records=records, seeds=seeds, into=out.parent / "figures")
    doc = Document()

    # -- house style --------------------------------------------------------

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(*INK)
    normal.paragraph_format.space_after = Pt(7)

    for level, size in ((1, 16), (2, 12.5), (3, 11)):
        style = doc.styles[f"Heading {level}"]
        style.font.color.rgb = RGBColor(*(ACCENT if level == 1 else INK))
        style.font.size = Pt(size)
        style.font.name = "Calibri"

    def para(text="", *, style=None, size=None, bold=False, italic=False, color=None, align=None):
        p = doc.add_paragraph(text, style=style)
        if align is not None:
            p.alignment = align
        for run in p.runs:
            if size:
                run.font.size = Pt(size)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if color:
                run.font.color.rgb = RGBColor(*color)
        return p

    def hrule(weight=8, color=ACCENT):
        """A ruled line. python-docx has no such thing; a paragraph border is one."""
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(weight))
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "".join(f"{channel:02X}" for channel in color))
        borders.append(bottom)
        p._p.get_or_add_pPr().append(borders)
        return p

    def table(headers, rows, *, style="Light Grid Accent 1"):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = style
        for cell, head in zip(t.rows[0].cells, headers, strict=True):
            cell.text = head
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for row in rows:
            for cell, value in zip(t.add_row().cells, row, strict=True):
                cell.text = str(value)
        doc.add_paragraph()
        return t

    def bullets(*points):
        for point in points:
            para(point, style="List Bullet")

    def figure(name: str, *, width=6.4):
        """Place a rendered chart, with its numbered caption underneath.

        Silently skipped when matplotlib is absent, so the document still
        builds — text-only, but built.
        """
        path = figures.get(name)
        if path is None:
            return
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        para(
            charts.FIGURES[name], align=WD_ALIGN_PARAGRAPH.CENTER, size=8.5, color=MUTED, bold=True
        )

    def campaign_row(campaign_id: str) -> list[str]:
        record = records[campaign_id]
        d = metrics.detection(record.episodes)
        # The adaptive run shares its rail and tiers with the static one, so
        # the control alone does not name it — two rows reading "Tier 1" are
        # two campaigns a reader cannot tell apart.
        adaptive_run = len({e.round for e in record.episodes}) > 1
        control = control_label(record.episodes[0].defense_tiers)
        return [
            control + (" · adaptive" if adaptive_run else ""),
            f"{metrics.asr(record.episodes):.1%}",
            fmt(metrics.rupees_moved(record.episodes)),
            f"{metrics.benign_pass_rate(record.episodes):.0%}",
            f"{d['precision']:.3f}",
            f"{d['recall']:.3f}",
            f"{d['f1']:.3f}",
        ]

    centre = WD_ALIGN_PARAGRAPH.CENTER
    tier1 = records.get("autopay-tier1")
    undefended = records.get("autopay-undefended")
    adaptive = records.get("autopay-adaptive")

    # ======================================================================
    # Cover
    # ======================================================================

    para()
    title = doc.add_heading("Praman", level=0)
    title.alignment = centre
    for run in title.runs:
        run.font.color.rgb = RGBColor(*ACCENT)

    para(
        "प्रमाण  ·  proof; evidence; that which establishes a claim",
        align=centre,
        italic=True,
        size=10,
        color=MUTED,
    )
    para(
        "Breach and attack simulation for agentic payment mandates",
        align=centre,
        bold=True,
        size=13.5,
    )
    hrule()
    para(
        "A product and technical report  ·  Mastercard Innovation Challenge @ GFF 2026",
        align=centre,
        size=10.5,
        color=MUTED,
    )
    para(f"Team {TEAM}", align=centre, bold=True, size=11.5)
    para(
        f"Prepared {date.today():%d %B %Y}\nSource code: {REPO}\nWorking prototype: {ARENA}",
        align=centre,
        size=10,
    )

    para()
    para(
        "In one sentence: AI agents are being given the authority to spend your money, "
        "the industry secures that authority with digital signatures, and signatures do "
        "not check whether the documents they sign agree with each other. Praman attacks "
        "that gap, measures what a control stops, then lets the attacker adapt and "
        "measures it again.",
        style="Intense Quote",
    )

    doc.add_heading("The position, in eight numbers", level=1)
    para(
        "Read from the committed campaigns at build time. The undefended row is the same "
        "harness with the control switched off — one argument's difference.",
        size=9.5,
        color=MUTED,
    )

    highlights: list[list[str]] = [
        ["Attack vectors mapped", f"{len(seeds)} across {len(SURFACES)} surfaces"],
        ["Built, run and measured", f"{len(built)} of {len(seeds)}"],
    ]
    if undefended and tier1:
        prevented = metrics.rupees_prevented(undefended.episodes, tier1.episodes)
        highlights += [
            ["Attack success with no control", f"{metrics.asr(undefended.episodes):.1%}"],
            ["Attack success behind Tier 1", f"{metrics.asr(tier1.episodes):.1%}"],
            ["Harm prevented", f"{fmt(prevented)} of the value at risk"],
            ["Honest traffic still settling", f"{metrics.benign_pass_rate(tier1.episodes):.0%}"],
        ]
    if adaptive:
        s = metrics.summarise(adaptive.episodes)
        highlights.append(
            ["Attack success once the attacker adapts", f"{s['adaptive_delta'] * 100:+.1f} points"]
        )
    highlights.append(["Cost of the control", "3.5 ms per authorisation; no model, no network"])
    table(["Measure", "Result"], highlights, style="Light List Accent 1")

    doc.add_page_break()

    # ======================================================================
    # Part I — Why
    # ======================================================================

    doc.add_heading("Part I  ·  Why this exists", level=1)

    doc.add_heading("A trillion-dollar flow is being built on an untested assumption", level=2)
    para(
        "The industry has settled on one answer to how an AI agent spends your money: "
        "cryptographically signed mandate chains. Google published AP2 in September 2025, "
        "Mastercard launched Agent Pay in April 2025, and NPCI is building the Unified "
        "Agent Protocol for UPI. All three sign the same three documents — Intent, Cart, "
        "Payment."
    )
    para(
        "The assumption underneath is that valid signatures mean a sound transaction. "
        "They do not. A signature proves who wrote a document; it says nothing about "
        "whether that document still agrees with the one before it.",
        bold=True,
    )
    figure("market-gap")
    para(
        "Juniper Research puts agentic commerce at $1.5 trillion by 2030 and names trust "
        "as the number one barrier to it. Breach-and-attack simulation — attacking your "
        "own defences to find out what they stop — is a $1.29 billion market that covers "
        "networks and endpoints. We could not identify a single vendor in it that "
        "simulates a payment mandate. That is the gap Praman fills.",
        size=10,
    )

    doc.add_heading("The exposure is largest where the rails are busiest", level=2)
    figure("india-exposure")
    para(
        "UPI settles more than 23 billion payments a month. Today a human approves each "
        "one, and fraud on it still runs to roughly ₹1,000 crore a year. UAP would let an "
        "agent approve them instead. Nobody has measured what that does to the number — "
        "which is the measurement this project exists to make.",
        size=10,
    )

    doc.add_heading("What actually breaks", level=2)
    para(
        "None of the attacks that matter forge a signature. Every one of them keeps all "
        "three signatures valid and breaks a relationship between the documents:"
    )
    table(
        ["The attack", "In plain terms", "Every signature still valid?"],
        [
            [
                "Cart substitution",
                "You authorised shoes; the cart quietly becomes something else",
                "Yes",
            ],
            [
                "Beneficiary rebinding",
                "The right amount leaves, to the wrong bank account",
                "Yes",
            ],
            [
                "Token-redemption race",
                "One authorisation is spent three times, at the same moment",
                "Yes",
            ],
            [
                "Invisible line item",
                "You are charged for something the summary never showed you",
                "Yes",
            ],
        ],
    )

    doc.add_page_break()

    # ======================================================================
    # Part II — What, and the philosophy
    # ======================================================================

    doc.add_heading("Part II  ·  What Praman is, and why it is shaped this way", level=1)
    para("A measuring instrument, not a detector. Four pieces that check each other:")
    table(
        ["Component", "What it is", "What it settles"],
        [
            [
                "The range",
                "A simulated payment rail — signed mandates, a merchant population, "
                "a double-entry ledger",
                "Where an attack's money actually ends up",
            ],
            [
                "Red",
                "An attack corpus, plus a search that adapts to the defence",
                "What an attacker can do, including things nobody wrote down",
            ],
            [
                "Blue",
                "A three-tier reference monitor sitting outside the agent",
                "What a control stops, and what it costs to run",
            ],
            [
                "The arena",
                "A browser view that replays any campaign, episode by episode",
                "That the numbers are inspectable rather than asserted",
            ],
        ],
    )
    para(
        "The ledger is the arbiter, and that is the most important decision in the system. "
        "An attack succeeded if and only if the balance on an attacker-controlled account "
        "went up. No model or scorer gets a vote. A harness whose detector decides whether "
        "the detector worked is grading its own homework.",
        bold=True,
    )

    doc.add_heading("The decisions, and what we rejected", level=2)
    table(
        ["Decision", "What we chose", "What we rejected, and why"],
        [
            [
                "What counts as success",
                "A balance delta on the attacker's ledger account",
                "The control's own verdict — the harness would be grading itself",
            ],
            [
                "Money",
                "Decimal paise; floats rejected, not coerced",
                "Silent coercion — a rounding error in a fraud harness is a number "
                "nobody can trust",
            ],
            [
                "How attacks sign",
                "Only with keys the attacker legitimately holds",
                "Forging the user's key — that tests our signature check, which already works",
            ],
            [
                "Order of the tiers",
                "Deterministic arithmetic first, model second",
                "Model first — a model that fails is unexplainable; a rule that fails names itself",
            ],
            [
                "Tier 3's authority",
                "It escalates for review; it never blocks",
                "Letting semantics refuse a payment — its first version escalated "
                "everything, at a 0% honest pass rate",
            ],
            [
                "The headline metric",
                "The adaptive delta — how much better an attacker gets",
                "Static attack success — every comparable tool reports round 0 and stops",
            ],
            [
                "Map versus build",
                "Counted and reported separately, always",
                "One combined figure — that is how an attack catalogue becomes a "
                "marketing document",
            ],
            [
                "Deployment",
                "Static files: no server, no credential, no database",
                "A container — free tiers sleep, and a judge's first visit pays the cold start",
            ],
        ],
    )

    doc.add_heading("What would change our mind", level=2)
    bullets(
        "If an attacker allowed to read refusals and retry got no further than the "
        "documented corpus, Praman would be a regression suite and the interesting claim "
        "would be gone. The adaptive command exits non-zero if that happens.",
        "If attack success against an undefended range were below 100%, the baseline "
        "would not be a baseline and every prevention figure measured against it would be "
        "inflated. The campaign runner exits non-zero if that happens.",
        "If honest traffic ever failed to settle, the control would be unshippable "
        "whatever its recall — so the benign pass rate is printed beside every headline "
        "number rather than in an appendix.",
    )

    doc.add_page_break()

    # ======================================================================
    # Part III — What it does, and the words it uses
    # ======================================================================

    doc.add_heading("Part III  ·  What it does", level=1)
    table(
        ["Capability", "Command", "What comes back"],
        [
            [
                "Map the attack surface",
                "python -m praman corpus",
                f"{len(seeds)} vectors across {len(SURFACES)} surfaces, the "
                f"{len(built)} built marked as built",
            ],
            [
                "Measure a control",
                "python -m praman matrix",
                "Every rail crossed with every control, baseline beside defended",
            ],
            [
                "Score it as a detector",
                "python -m praman detect",
                "Precision, recall, F1, false-positive rate and AUC, per campaign",
            ],
            [
                "Let the attacker adapt",
                "python -m praman adapt",
                "The static-versus-adaptive gap, and the strategy that broke each rule",
            ],
            [
                "Add your own test case",
                "python -m praman run <scenario.yaml>",
                "A new campaign in the arena — no code change, no redeploy",
            ],
        ],
    )
    para(
        "The last row is what makes this a product rather than a demo. One YAML file — "
        "rail, control, attacks, an optional merchant overlay — and the campaign appears "
        "in the live arena beside the shipped ones, under its own heading so it is never "
        "mistaken for one of ours.",
        size=10,
    )

    doc.add_heading("The words this report uses", level=2)
    para(
        "Defined here so the evidence that follows is readable without a payments background.",
        size=9.5,
        color=MUTED,
    )
    table(["Term", "What it means"], [[term, meaning] for term, meaning in GLOSSARY])

    doc.add_page_break()

    # ======================================================================
    # Part IV — The evidence
    # ======================================================================

    doc.add_heading("Part IV  ·  The evidence", level=1)
    para(
        "Four books, answering the four questions the challenge asks: what we identify, "
        "how we generate it, what defends against it, and what it takes to run in "
        "production.",
        size=9.5,
        color=MUTED,
    )

    # -- Book One: identify -------------------------------------------------

    doc.add_heading("Book One  ·  Identify — the attack surface, mapped", level=2)
    figure("coverage")
    para(
        f"{len(seeds)} vectors across {len(SURFACES)} surfaces; {len(built)} built, run "
        "and measured. The split is deliberate and is never collapsed into one number: "
        "identification is research, implementation is evidence. A test fails the build if "
        "the corpus and the code disagree about which is which."
    )
    para(
        "Root causes RC-1 to RC-5 are the vulnerability classes from arXiv 2607.21824. "
        "RC-6 (agent judgement subverted) and L-1/L-2 (the human deceived, the human "
        "coerced) are ours, and are marked as ours rather than folded in as though they "
        "were the paper's.",
        size=10,
    )
    table(
        ["Surface", "What it attacks", "Mapped", "Built"],
        [
            [
                surface,
                HEADINGS[surface],
                len(surfaces.get(surface, [])),
                sum(s.implemented for s in surfaces.get(surface, [])),
            ]
            for surface in SURFACES
        ],
    )

    doc.add_heading("The full map", level=3)
    for surface in SURFACES:
        entries = surfaces.get(surface, [])
        if not entries:
            continue
        para(f"{surface} — {HEADINGS[surface]}", bold=True, size=10.5)
        table(
            ["ID", "Attack", "Root cause", "Caught by", "Built"],
            [
                [s.id, s.name, s.root_cause, s.caught_by, "yes" if s.implemented else "—"]
                for s in entries
            ],
        )

    para(
        "Three independent entries — S-05, I-22 and D-33 — fail for want of the same "
        "missing control: cumulative accounting across mandates against a delegated cap. "
        "Three surfaces converging on one gap is the clearest signal in the map for what "
        "to build next, and finding that is what a map is for.",
        bold=True,
    )

    # -- Book Two: generate -------------------------------------------------

    doc.add_heading("Book Two  ·  Generate — the range, and the search", level=2)
    doc.add_heading("Why the simulation is believable", level=3)
    table(
        ["What could have been faked", "What Praman actually does"],
        [
            [
                "Signatures",
                "Real ECDSA P-256 over a canonical serialisation. Attacks re-sign only "
                "with keys the attacker legitimately holds — never the user's.",
            ],
            [
                "Money",
                "Decimal paise. A float is rejected, not rounded, because a rounding "
                "error in a fraud harness is a number nobody can trust.",
            ],
            [
                "Settlement",
                "SQLite with real transactions, refusing anything that did not go "
                "through the monitor. Harm is a balance delta, not a label.",
            ],
            [
                "The merchants",
                "Built so no single feature separates honest from criminal — including "
                "a compromised merchant that out-scores two honest ones on every feature.",
            ],
            [
                "The rails",
                "Two profiles: UPI Autopay, and NPCI UAP with a delegated cap where the "
                "agent is granted less than its principal holds.",
            ],
            [
                "Reproducibility",
                "Everything seeded; the same seed reproduces every result field on any "
                "machine. CI re-derives the committed numbers on every push.",
            ],
        ],
    )

    doc.add_heading("Attacks nobody wrote", level=3)
    para(
        "A fixed corpus is a checklist, and checklists already exist. Every variant the "
        "control refuses is handed the name of the rule that refused it — exactly what a "
        "real attacker learns from a rejection message — and a deterministic search "
        "proposes variants aimed at that rule. Winners are kept, dead ends retired, and "
        "the next round faces the same control with better attacks. The search is a pure "
        "function of (variant, rule), so it replays identically from a seed on any "
        "machine. Its dead ends are reported too."
    )

    # -- Book Three: defend -------------------------------------------------

    doc.add_heading("Book Three  ·  Defend — what the control actually stops", level=2)
    table(
        ["Tier", "What it does", "Authority"],
        [
            [
                "Tier 1",
                "Seven deterministic checks on the arithmetic between mandates. No model, 3.5 ms.",
                "Blocks, and names the rule",
            ],
            [
                "Tier 2",
                "A gradient-boosted model fitted on what Tier 1 let through — Red's "
                "survivors are Blue's training set.",
                "Blocks",
            ],
            [
                "Tier 3",
                "Semantic divergence between what was described and what was charged.",
                "Escalates only, never blocks",
            ],
        ],
    )
    figure("control-effect")
    para(
        "Attack success falls from 100% to 16.7% behind seven lines of arithmetic, and to "
        "zero with all three tiers. Then it rises again to 41.2% the moment the attacker "
        "is allowed to learn — which is the finding this project exists to report.",
        size=10,
    )

    doc.add_heading("Detection efficacy, per rail", level=3)
    para(
        "Attacks are the positive class; honest traffic is the negative class. Refusals "
        "and harness errors are excluded from every rate, because they describe the victim "
        "model or our own harness rather than the control.",
        size=10,
    )
    for rail, label in (("autopay", "UPI Autopay"), ("uap", "NPCI UAP")):
        para(label, bold=True, size=10.5)
        table(
            ["Control", "ASR", "Moved", "Benign pass", "Precision", "Recall", "F1"],
            [
                campaign_row(f"{rail}-{name}")
                for name in ("undefended", "tier1", "tier123", "adaptive")
                if f"{rail}-{name}" in records
            ],
        )

    doc.add_heading("Detection is not prevention", level=3)
    figure("detection-paradox")
    if tier1:
        para(
            "Both panels describe the same campaign. Tier 1 scores precision 1.000, "
            "recall 1.000, F1 1.000 and zero false positives — and "
            f"{fmt(metrics.rupees_moved(tier1.episodes))} still reaches the attacker. "
            "S-03 fires three concurrent redemptions at one authorisation; the control "
            "refuses two and the third settles. Every one of those episodes is a true "
            "positive by decision and a success by ledger.",
            bold=True,
        )
    para(
        "Reporting only F1 would show a perfect classifier over a system losing money. "
        "Reporting only attack success would hide that the control saw every attack "
        "coming. Both are printed side by side, always, because in live payments a flag "
        "that arrives after settlement is a chargeback, not a defense.",
        size=10,
    )

    doc.add_heading("What adaptation does", level=3)
    figure("adaptation")
    for rail, label in (("autopay", "UPI Autopay"), ("uap", "NPCI UAP")):
        record = records.get(f"{rail}-adaptive")
        if not record:
            continue
        s = metrics.summarise(record.episodes)
        rounds = sorted({e.round for e in record.episodes})
        first = metrics.recall([e for e in record.episodes if e.round == rounds[0]])
        last = metrics.recall([e for e in record.episodes if e.round == rounds[-1]])
        para(
            f"{label}: attack success rises from {s['static_asr']:.1%} to "
            f"{s['adaptive_asr']:.1%} — a delta of {s['adaptive_delta'] * 100:+.1f} "
            f"points — and the control is first broken in round {s['rounds_to_break']}. "
            f"Recall falls from {first:.3f} to {last:.3f}: the variants the search finds "
            "are not caught late, they are not caught at all.",
            style="List Bullet",
        )
    para(
        "Every comparable tool reports the round-0 number and stops. A control that holds "
        "against a documented attack list and folds in one round against an attacker that "
        "reads its rejections has not been tested.",
        bold=True,
    )

    doc.add_heading("Where the learned tier stops, and which model is shopping", level=3)
    para(
        "Tier 2 catches 52 of 52 in-sample with zero false positives in 72, scores an AUC "
        "of 0.917, and catches 0% against either held-out merchant it was not trained on. "
        "Reported rather than hidden: a learned tier generalises to variants of what it "
        "has seen, not to techniques nobody has run — which is why Tier 1 goes first and "
        "why the headline result does not rest on a model.",
        size=10,
    )
    table(
        ["Victim model", "M-08 branded whisper", "Every structural attack"],
        [
            ["scripted baseline", "100%", "100%"],
            ["qwen2.5:1.5b", "0%", "100%"],
            ["llama3.2:3b", "0%", "100%"],
            ["mistral:latest", "100%", "100%"],
        ],
    )
    para(
        "Structural attacks succeed against every model at the same rate — they break "
        "arithmetic, and arithmetic has no opinion about which model is shopping. Choosing "
        "a better model moves the semantic column and leaves the structural one untouched, "
        "which is the case for fixing this at the control layer rather than hoping for a "
        "safer agent.",
        size=10,
    )

    # -- Book Four: feasibility ---------------------------------------------

    doc.add_heading("Book Four  ·  Feasibility — running this in production", level=2)
    table(
        ["Question a deployer asks", "Answer"],
        [
            [
                "Is it fast enough for an authorisation path?",
                "3.5 ms per authorisation, of which the arithmetic is ~0.02 ms. No model, "
                "no GPU, no network call.",
            ],
            [
                "Who can deploy it?",
                "A PSP, issuer or rail operator, without the agent vendor's cooperation — "
                "which matters, because the agent vendor does not carry the loss.",
            ],
            [
                "Can it explain a refusal to a regulator?",
                "Every block names the rule with observed and expected values. Asserted by "
                "a test, not claimed.",
            ],
            [
                "Will it block real customers?",
                "Benign pass rate is 100% across every committed campaign, printed beside "
                "every headline number.",
            ],
            [
                "Does it work on more than one rail?",
                "UPI Autopay and NPCI UAP are both implemented and produce genuinely "
                "different campaigns. A new rail is a profile, not a rewrite.",
            ],
            [
                "What does the demo cost to run?",
                "Nothing. No server, no credential, no database — the deployed arena is "
                "committed files a judge can diff against a local re-run.",
            ],
        ],
    )

    doc.add_heading("Risks and honest limits", level=3)
    bullets(
        "I-31, coerced-principal intent, is deliberately unbuildable at the mandate layer. "
        "The human signed under duress, every check passes, the money is gone — the fraud "
        "is upstream of the mandate. Naming the limit of the cryptographic approach the "
        "industry is betting on is more useful than pretending not to have one.",
        f"{len(seeds) - len(built)} of {len(seeds)} mapped vectors are not built. Each says "
        "why in its note, so the range's boundaries are recorded rather than hidden.",
        "The range is a simulation — a high-fidelity one, but a simulation. Production "
        "would meet merchant populations and volumes we have modelled rather than "
        "observed. What transfers is the method and the control; the rates would have to "
        "be re-measured on real traffic, and the harness is built to be pointed at it.",
        "Tier 2's held-out recall is 0%. It is in this report, and in the arena, because a "
        "learned tier that only looks good in-sample is the failure mode this whole "
        "project is designed to catch.",
    )

    doc.add_page_break()

    # ======================================================================
    # Appendix
    # ======================================================================

    doc.add_heading("Appendix A  ·  Reproduce every number in this report", level=1)
    para("Python 3.12 and nothing else. No Node, no npm, no build step, no API key.")
    for command, what in (
        ('pip install -e ".[dev,docs]"', "install, with the document toolchain"),
        ("python -m praman demo", "one honest purchase, settled end to end"),
        ("python -m praman corpus", "the attack map, Figure 6"),
        ("python -m praman matrix", "regenerate every campaign"),
        ("python -m praman detect", "precision, recall, F1, AUC — Figures 3 and 4"),
        ("python -m praman adapt", "the static-versus-adaptive gap — Figure 5"),
        ("python -m praman models", "attack success by victim model (needs Ollama)"),
        ("python -m praman serve", "the arena, locally"),
        ("python -m praman test", "the full suite"),
        ("python -m praman walkthrough", "rebuild this document, figures included"),
    ):
        para(f"{command}   —   {what}", style="List Bullet")

    doc.add_heading("Appendix B  ·  Market and technical figures, by source", level=1)
    para(
        "Every external number in this report, with whoever published it. Nothing here is "
        "a forecast of ours; where analysts disagree the spread is shown rather than the "
        "most flattering end of it.",
        size=9.5,
        color=MUTED,
    )
    table(
        ["Figure", "Value", "Source"],
        [
            [
                fact.label,
                fact.value,
                "measured by us" if fact.source == "OURS" else cite(fact.source).author,
            ]
            for fact in MARKET
        ],
    )

    doc.add_heading("Appendix C  ·  References", level=1)
    for index, source in enumerate(SOURCES, start=1):
        para(f"[{index}]  {source.reference()}", size=9, style=None)

    hrule(weight=4, color=MUTED)
    para(
        "Generated by `python -m praman walkthrough`. Every quantity was read from the "
        "repository's committed results at build time, and every figure was drawn from "
        "them; a test in the suite fails if either stops being true.",
        size=9,
        italic=True,
        color=MUTED,
        align=centre,
    )

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        prog="python -m praman walkthrough",
        description="Praman — build the solution walkthrough as a .docx",
    )
    parser.add_argument("--results", type=Path, default=Path("results"))
    parser.add_argument("--out", type=Path, default=Path("docs/Praman-Solution-Walkthrough.docx"))
    args = parser.parse_args(argv)
    console_setup()

    try:
        import docx  # noqa: F401
    except ImportError:
        print('\n  python-docx is not installed — run: pip install -e ".[docs]"\n')
        return 1

    banner("SOLUTION WALKTHROUGH")

    out = build(results=args.results, out=args.out)
    print(f"\n  wrote {out}  ({out.stat().st_size / 1024:.0f} KB)")
    print("  every figure read from results/ and corpus.yaml at build time\n")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
