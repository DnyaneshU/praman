"""The submission walkthrough — a required artifact, so it is built by a test.

The challenge asks for a Word document covering the attacks identified, how
they are simulated, the detection model with its efficacy, and real-world
feasibility. It is generated from `results/` and `corpus.yaml` rather than
typed, so it cannot drift from the repository it describes — and these assert
that it still builds and still says what the harness says.

A hand-written walkthrough fails silently: the numbers go stale and nobody
notices until a judge re-runs the code. A generated one fails loudly, here.
"""

from __future__ import annotations

import importlib.util

import pytest

from praman import metrics
from praman.api.replay import CampaignStore
from praman.red.corpus import SURFACES, load_corpus

needs_docx = pytest.mark.skipif(
    importlib.util.find_spec("docx") is None,
    reason='python-docx is not installed (pip install -e ".[docs]")',
)


@pytest.fixture(scope="module")
def document(tmp_path_factory):
    from praman.report.walkthrough import build

    out = tmp_path_factory.mktemp("walkthrough") / "walkthrough.docx"
    from docx import Document

    return Document(build(results="results", out=out))


@needs_docx
def test_it_covers_all_four_things_the_challenge_asks_for(document):
    headings = " ".join(
        p.text.lower() for p in document.paragraphs if p.style.name.startswith("Heading")
    )
    for required in ("identify", "generate", "defend", "feasibility"):
        assert required in headings, f"no section covering {required}"


@needs_docx
def test_the_attack_map_is_complete(document):
    """Every surface in the corpus reaches the document."""
    text = " ".join(p.text for p in document.paragraphs)
    for surface in SURFACES:
        assert surface in text, f"{surface} is mapped but absent from the walkthrough"


@needs_docx
def test_every_mapped_attack_appears(document):
    """A map that lists 34 vectors and prints 12 would be a claim, not a map."""
    printed = {cell.text for table in document.tables for row in table.rows for cell in row.cells}
    for seed in load_corpus():
        assert seed.id in printed, f"{seed.id} is in the corpus but not in the document"


@needs_docx
def test_the_numbers_are_the_harnesss_numbers(document):
    """The point of generating it: the document cannot disagree with results/.

    Checked against the campaign the whole submission rests on, so a re-run
    that changes the result changes the document or fails this test.
    """
    record = next(r for r in CampaignStore("results").list() if r.id == "autopay-adaptive")
    summary = metrics.summarise(record.episodes)
    text = " ".join(p.text for p in document.paragraphs)

    assert f"{summary['static_asr']:.1%}" in text
    assert f"{summary['adaptive_asr']:.1%}" in text
    assert f"{summary['adaptive_delta'] * 100:+.1f} points" in text


@needs_docx
def test_the_static_and_adaptive_runs_are_distinguishable(document):
    """Two rows reading "Tier 1" are two campaigns a reader cannot tell apart."""
    # Per table: there is one per rail, so a control legitimately repeats
    # across them. Within a rail it must not.
    tables = [t for t in document.tables if t.rows[0].cells[0].text == "Control"]
    assert tables, "no detection efficacy table in the document"
    for table in tables:
        controls = [row.cells[0].text for row in table.rows[1:]]
        assert len(controls) == len(set(controls)), f"indistinguishable rows: {controls}"


@needs_docx
def test_the_honest_limits_are_stated(document):
    """The document must carry the unflattering numbers too.

    Tier 2's held-out recall and the coerced-principal limit are the two places
    the submission admits a boundary. A walkthrough that dropped them would be
    a better-looking document and a worse one.
    """
    text = " ".join(p.text for p in document.paragraphs).lower()
    assert "held-out" in text
    assert "coerced-principal" in text
    assert "0%" in text


# -- the report as a report -------------------------------------------------
#
# The document is addressed to a reader deciding whether this is a product,
# not only to a rubric. These guard the shape that makes it one: why before
# what, what before how, and the reasoning shown rather than the conclusion
# asserted.


@needs_docx
def test_it_argues_why_before_it_explains_how(document):
    """A reader should reach the reasoning before the first metric.

    The engineering is Part IV on purpose. A walkthrough that opens on tables
    tells you what was measured and never what was at stake.
    """
    headings = [
        p.text for p in document.paragraphs if p.style.name.startswith("Heading") and p.text.strip()
    ]
    order = [next(i for i, h in enumerate(headings) if part in h) for part in ("Part I", "Part IV")]
    assert order == sorted(order), "the engineering precedes the argument for it"

    joined = " ".join(headings)
    for part in ("Part I", "Part II", "Part III", "Part IV", "Appendix"):
        assert part in joined, f"{part} is missing"


@needs_docx
def test_the_rejected_alternatives_are_named(document):
    """The decisions table is the philosophy, and half of it is what we did not do.

    A table of choices with no alternatives is a feature list. The rejected
    column is the part that shows the reasoning, so it has to carry reasons.
    """
    decisions = next(
        (t for t in document.tables if t.rows[0].cells[0].text == "Decision"),
        None,
    )
    assert decisions is not None, "no decisions table in the document"
    assert len(decisions.rows) > 5, "too few decisions to be an account of the design"

    for row in decisions.rows[1:]:
        decision, chosen, rejected = (c.text.strip() for c in row.cells)
        assert chosen, f"{decision} names no choice"
        # A rejected alternative has to carry the reasoning, not just a label —
        # "a container" explains nothing; "a container, because free tiers
        # sleep" is the part a reader learns from.
        assert len(rejected.split()) >= 8, f"{decision}: the rejected option gives no reason"


@needs_docx
def test_the_claim_is_falsifiable(document):
    """What would have to be true for this to be wrong, stated in the document.

    And wired to something: a falsifying condition nothing checks is a
    paragraph. Two of the three are exit codes the CLI actually returns.
    """
    headings = [p.text.lower() for p in document.paragraphs if p.style.name.startswith("Heading")]
    assert any("change our mind" in h for h in headings), "no falsifiability section"

    text = " ".join(p.text for p in document.paragraphs).lower()
    assert "exits non-zero" in text, "no falsifying condition is wired to a check"


@needs_docx
def test_the_opening_figures_are_the_harnesss_figures(document):
    """The summary table a skimmer reads first must agree with results/."""
    figures = next((t for t in document.tables if t.rows[0].cells[0].text == "Measure"), None)
    assert figures is not None, "no opening figures table"

    values = {row.cells[0].text: row.cells[1].text for row in figures.rows[1:]}
    store = {r.id: r for r in CampaignStore("results").list()}
    tier1 = store["autopay-tier1"]

    assert values["Attack success behind Tier 1"] == f"{metrics.asr(tier1.episodes):.1%}"
    assert (
        values["Attack success with no control"]
        == f"{metrics.asr(store['autopay-undefended'].episodes):.1%}"
    )
    assert str(len(load_corpus())) in values["Attack vectors mapped"]


@needs_docx
def test_the_cover_names_the_team(document):
    """It lives in the generator, so a rebuild cannot drop it.

    Typed into the .docx instead, it would survive exactly until the next
    `python -m praman walkthrough` — which is a command the document itself
    tells the reader to run.
    """
    from praman.report.walkthrough import TEAM

    cover = " ".join(p.text for p in document.paragraphs[:12])
    assert TEAM in cover, "the cover does not name the team"


# -- the figures ------------------------------------------------------------


needs_charts = pytest.mark.skipif(
    importlib.util.find_spec("matplotlib") is None,
    reason='matplotlib is not installed (pip install -e ".[docs]")',
)


@needs_docx
@needs_charts
def test_every_figure_reaches_the_document(document):
    """Six charts declared, six embedded.

    A figure that renders but is never placed is worse than no figure: the
    caption numbering skips and the reader goes looking for it.
    """
    from praman.report.charts import FIGURES

    embedded = [rel for rel in document.part.rels.values() if "image" in rel.reltype]
    assert len(embedded) == len(FIGURES), (
        f"{len(FIGURES)} figures declared, {len(embedded)} embedded"
    )

    captions = " ".join(p.text for p in document.paragraphs)
    for caption in FIGURES.values():
        assert caption in captions, f"no caption for {caption!r}"


@needs_docx
def test_the_document_stays_readable_rather_than_exhaustive(document):
    """Prose is capped so the figures and tables carry the argument.

    An earlier draft ran to nearly four thousand words of unbroken text. The
    brief was a report a non-specialist can read, and past roughly three
    thousand words that stops being true no matter how good the sentences are.
    """
    words = sum(len(p.text.split()) for p in document.paragraphs)
    assert words < 3000, f"{words} words — the prose is crowding out the evidence"
    assert len(document.tables) >= 15, "too few tables for a report this length"


@needs_docx
def test_the_jargon_is_defined_before_it_is_used(document):
    """A glossary, and it has to precede the engineering that needs it."""
    from praman.report.walkthrough import GLOSSARY

    paragraphs = [p.text for p in document.paragraphs]
    defined = {cell.text for t in document.tables for row in t.rows for cell in row.cells}
    for term, _ in GLOSSARY:
        assert term in defined, f"{term} is used but never defined"

    glossary_at = next(i for i, t in enumerate(paragraphs) if "words this report uses" in t.lower())
    evidence_at = next(i for i, t in enumerate(paragraphs) if t.startswith("Part IV"))
    assert glossary_at < evidence_at, "the glossary comes after the terms it defines"


@needs_docx
def test_external_numbers_are_attributed_in_the_document(document):
    """Appendix B exists and names a publisher for every borrowed figure."""
    from praman.report.sources import MARKET

    attributions = next(
        (t for t in document.tables if t.rows[0].cells[0].text == "Figure"),
        None,
    )
    assert attributions is not None, "no source table in the document"
    assert len(attributions.rows) - 1 == len(MARKET)
    for row in attributions.rows[1:]:
        assert row.cells[2].text.strip(), f"{row.cells[0].text} is unattributed"


@needs_docx
def test_the_references_are_numbered_and_resolvable(document):
    from praman.report.sources import SOURCES

    text = " ".join(p.text for p in document.paragraphs)
    for index, source in enumerate(SOURCES, start=1):
        assert f"[{index}]" in text, f"reference {index} is missing"
        assert source.url in text, f"{source.key} has no URL in the reference list"
